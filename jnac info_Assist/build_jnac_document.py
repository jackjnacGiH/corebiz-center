from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from html import escape


DATA = Path("jnac_site_data.json")
OUT = Path("JNAC_website_knowledge_base_final.docx")
OUT_DOC = Path("JNAC_website_knowledge_base_final.doc")


MAIN_PAGE_LABELS = {
    "เกี่ยวกับเรา",
    "ติดต่อเรา",
    "Catalogue",
    "วิธีการสั่งซื้อ",
    "บทความ",
}


def clean(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def page_title(page: dict) -> str:
    headings = [clean(h) for h in page.get("headings", []) if clean(h)]
    if headings:
        return headings[0]
    title = clean(page.get("title", ""))
    path = unquote(urlparse(page["url"]).path.strip("/").split("/")[-1])
    return title or path or page["url"]


def text_lines(page: dict) -> list[str]:
    return [clean(line) for line in page.get("text", "").splitlines() if clean(line)]


def core_lines(page: dict, limit: int = 12) -> list[str]:
    lines = text_lines(page)
    title = page_title(page)
    hit_indexes = [i for i, line in enumerate(lines) if title and title in line]
    start = hit_indexes[-1] if hit_indexes else 0
    end_candidates = [i for i in range(start + 1, len(lines)) if lines[i].lower() == "tweet"]
    end = end_candidates[0] if end_candidates else min(start + 18, len(lines))
    body = []
    for line in lines[start:end]:
        if line in {">", "ย้อนกลับ", "หน้าแรก", "Register", "Mail Box"}:
            continue
        if line not in body:
            body.append(line)
    return body[:limit]


def summary_text(page: dict, max_chars: int = 420) -> str:
    body = core_lines(page, limit=10)
    title = page_title(page)
    body = [line for line in body if line != title]
    text = " ".join(body)
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > max_chars:
        text = text[: max_chars - 3].rstrip() + "..."
    return text


def category_for(page: dict) -> str:
    headings = [clean(h) for h in page.get("headings", []) if clean(h)]
    if not headings:
        return "อื่น ๆ"
    if headings[0] in MAIN_PAGE_LABELS:
        return "ข้อมูลเว็บไซต์/บริษัท"
    if "บทความ" in page["url"] or any("บทความ" in h for h in headings):
        return "บทความ/ความรู้"
    if len(headings) > 1 and headings[1] not in {"สอบถาม สั่งซื้อสินค้า..", headings[0]}:
        return headings[1]
    return "สินค้า/หมวดสินค้า"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Tahoma"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(4.2), Cm(12.8)]
    for row_idx, (key, value) in enumerate(rows):
        row = table.rows[0] if row_idx == 0 else table.add_row()
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        set_cell_text(row.cells[0], key, bold=True, size=9)
        set_cell_text(row.cells[1], value, size=9)
        set_cell_shading(row.cells[0], "F2F5F7")
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float], font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        header.cells[idx].width = Cm(widths_cm[idx])
        set_cell_text(header.cells[idx], text, bold=True, size=font_size)
        set_cell_shading(header.cells[idx], "E8EEF3")
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            row.cells[idx].width = Cm(widths_cm[idx])
            set_cell_text(row.cells[idx], text, size=font_size)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Tahoma"
    run.font.size = Pt(10.5)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Tahoma"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(55, 86, 35)
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True

    header = section.header.paragraphs[0]
    header.text = "JNAC Website Knowledge Base"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Tahoma"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(110, 110, 110)

    footer = section.footer.paragraphs[0]
    footer.text = "Source: www.jnac.co.th"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Tahoma"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(110, 110, 110)


def main() -> None:
    data = json.loads(DATA.read_text(encoding="utf-8"))
    pages = data["pages"]
    doc = Document()
    setup_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("เอกสารฐานข้อมูลเว็บไซต์ JNAC สำหรับสร้างแชทบอท")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("รวบรวมจากเว็บไซต์ www.jnac.co.th เพื่อใช้เป็น knowledge base ของธุรกิจ บริษัท เจ แนค (ประเทศไทย) จำกัด")

    add_kv_table(
        doc,
        [
            ("แหล่งข้อมูล", data["source"]),
            ("วันที่เก็บข้อมูล", data["fetched_at"]),
            ("จำนวนหน้าที่ดึงข้อมูล", f"{data['page_count']} หน้า"),
            ("ขอบเขต", "หน้าบริษัท ติดต่อ วิธีสั่งซื้อ Catalogue บทความ และหน้าสินค้า/หมวดสินค้าภายในโดเมน"),
        ],
    )

    doc.add_heading("1. ภาพรวมธุรกิจ", level=1)
    add_bullet(doc, "ชื่อบริษัท: บริษัท เจ แนค (ประเทศไทย) จำกัด / J NAC (THAILAND) CO., LTD.")
    add_bullet(doc, "ธุรกิจหลัก: ผลิต จำหน่าย นำเข้า ส่งออก และรับจ้างผลิตสินค้าในกลุ่มอุตสาหกรรมงานขัด ปัด ตัด เจียร รวมถึงงาน ICT Fixture, ICT Tester, Function Test, Assembly Part & Machine")
    add_bullet(doc, "สินค้าหลักจากเว็บไซต์: กระดาษทราย จานทรายซ้อน ใบเจียร ใบตัด ล้อทราย ใยขัดสก๊อตไบร์ท หินขัด ลูกขัด เครื่องมือขัด เครื่องมือลม และสินค้างานขัดแม่พิมพ์")
    add_bullet(doc, "จุดยืน: แหล่งรวมสินค้าสำหรับงานขัด ตัด เจาะ เจียร และงานอุตสาหกรรม พร้อมประสบการณ์ด้านการผลิตและจัดจำหน่าย")

    doc.add_heading("2. ข้อมูลติดต่อและบริษัท", level=1)
    add_kv_table(
        doc,
        [
            ("ชื่อภาษาไทย", "บริษัท เจ แนค (ประเทศไทย) จำกัด"),
            ("ชื่อภาษาอังกฤษ", "J NAC (THAILAND) CO., LTD."),
            ("ที่อยู่", "84 หมู่ 2 ซ.สุนทรวิภาค ถ.บางพลี-ตำหรุ ต.แพรกษาใหม่ อ.เมือง จ.สมุทรปราการ 10280"),
            ("Address", "84 Moo.2 Soi.Sunthornviphakh, Bangplee-tamru Rd., Prakasamai, Muang, Samutprakarn 10280"),
            ("โทรศัพท์", "0 2101 5587"),
            ("มือถือ", "08 0016 1700"),
            ("แฟกซ์", "0 2101 5587"),
            ("อีเมล", "info@jnac.co.th, jnac.co.th@gmail.com"),
            ("เว็บไซต์", "www.jnac.co.th"),
            ("Line ID", "@jnac"),
            ("เลขประจำตัวผู้เสียภาษี", "0115561012346"),
        ],
    )

    doc.add_heading("3. วิธีการสั่งซื้อและชำระเงิน", level=1)
    ordering = next((p for p in pages if page_title(p) == "วิธีการสั่งซื้อ"), None)
    if ordering:
        for line in core_lines(ordering, limit=24)[1:]:
            add_bullet(doc, line)

    doc.add_heading("4. หมวดหมู่สินค้าและความรู้จากเว็บไซต์", level=1)
    categories = Counter(category_for(page) for page in pages)
    category_rows = [[cat, str(count)] for cat, count in categories.most_common()]
    add_table(doc, ["หมวด", "จำนวนหน้า"], category_rows, [13.5, 3.5], font_size=9)

    doc.add_heading("5. รายการหน้าข้อมูลสำหรับทำแชทบอท", level=1)
    doc.add_paragraph(
        "ตารางนี้เป็นดัชนีความรู้หลักของเว็บไซต์ โดยตัดเมนูและส่วนท้ายซ้ำออก เหลือหัวข้อ หมวด เนื้อหาสำคัญ และ URL ต้นทางสำหรับอ้างอิง/ทำ retrieval."
    )
    knowledge_rows = []
    for idx, page in enumerate(pages, start=1):
        knowledge_rows.append(
            [
                str(idx),
                page_title(page),
                category_for(page),
                summary_text(page),
                page["url"],
            ]
        )
    add_table(doc, ["#", "หัวข้อ/สินค้า", "หมวด", "ข้อมูลสำคัญที่ดึงได้", "URL"], knowledge_rows, [1.0, 4.0, 3.0, 6.3, 2.7], font_size=6)

    doc.add_heading("6. บทความและหัวข้อความรู้", level=1)
    article_pages = [
        p
        for p in pages
        if p["url"].split("/")[3:4] and (p["url"].split("/")[3] in {"17002546", "17002573", "17003398", "17004068", "17005754", "17015052", "17017795", "17428113"} or "บทความ" in page_title(p))
    ]
    article_rows = [[page_title(page), summary_text(page, 520), page["url"]] for page in article_pages]
    add_table(doc, ["หัวข้อบทความ", "สาระสำคัญ", "URL"], article_rows, [4.8, 9.0, 3.2], font_size=7)

    doc.add_heading("7. แนวทางแปลงเป็นฐานข้อมูลแชทบอท", level=1)
    add_bullet(doc, "ใช้ `หัวข้อ/สินค้า` เป็นชื่อ document หรือ intent candidate และใช้ `หมวด` เป็น metadata สำหรับ filtering")
    add_bullet(doc, "ใช้ `ข้อมูลสำคัญที่ดึงได้` เป็นข้อความ chunk เริ่มต้น และเก็บ `URL` เป็น source citation")
    add_bullet(doc, "คำถามติดต่อ/สั่งซื้อควรตอบจากส่วนข้อมูลติดต่อและวิธีการสั่งซื้อเท่านั้น เพราะเป็นข้อมูลเชิงนโยบายและธุรกรรม")
    add_bullet(doc, "คำถามสินค้าให้ตอบชื่อสินค้า หมวด คุณสมบัติ/การใช้งานที่ปรากฏ และปิดท้ายด้วยช่องทางสอบถามเมื่อไม่มีราคา/สต๊อกบนหน้าเว็บ")
    add_bullet(doc, "หากทำระบบ RAG ควรแยก chunk ตามแถวสินค้าและบทความ ไม่ควรนำเมนูซ้ำเข้าฐานข้อมูล")

    doc.add_heading("8. หมายเหตุคุณภาพข้อมูล", level=1)
    add_bullet(doc, "เว็บไซต์มีเมนูสินค้าและส่วนท้ายซ้ำในทุกหน้า เอกสารนี้สกัดเฉพาะช่วงเนื้อหาหลักต่อหน้าเพื่อลด noise")
    add_bullet(doc, "บางหน้ามีข้อมูลรายละเอียดสั้นหรือเป็นหน้าหมวด จึงเหมาะใช้เป็นดัชนีสินค้าและ metadata มากกว่าข้อความตอบยาว")
    add_bullet(doc, "ราคาและสต๊อกไม่พบเป็นข้อมูลมาตรฐานในหน้าที่เก็บ จึงควรให้แชทบอทตอบให้ติดต่อบริษัทเพื่อยืนยันราคา/จำนวนคงเหลือ")

    doc.core_properties.title = "JNAC Website Knowledge Base"
    doc.core_properties.subject = "Structured website data for chatbot knowledge base"
    doc.core_properties.author = "Codex"
    doc.core_properties.created = datetime.now()
    doc.save(OUT)
    write_html_doc(data, pages)
    print(OUT.resolve())
    print(OUT_DOC.resolve())


def write_html_doc(data: dict, pages: list[dict]) -> None:
    rows = []
    for idx, page in enumerate(pages, start=1):
        rows.append(
            "<tr>"
            f"<td>{idx}</td>"
            f"<td>{escape(page_title(page))}</td>"
            f"<td>{escape(category_for(page))}</td>"
            f"<td>{escape(summary_text(page))}</td>"
            f"<td>{escape(page['url'])}</td>"
            "</tr>"
        )
    category_rows = "\n".join(
        f"<tr><td>{escape(cat)}</td><td>{count}</td></tr>"
        for cat, count in Counter(category_for(page) for page in pages).most_common()
    )
    html = f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>JNAC Website Knowledge Base</title>
<style>
body {{ font-family: Tahoma, Arial, sans-serif; font-size: 11pt; color: #222; }}
h1 {{ color: #1f4e79; font-size: 22pt; }}
h2 {{ color: #375623; font-size: 15pt; margin-top: 24px; }}
table {{ border-collapse: collapse; width: 100%; margin: 10px 0 20px; }}
th, td {{ border: 1px solid #b9c3cc; padding: 6px; vertical-align: top; }}
th {{ background: #e8eef3; }}
.meta td:first-child {{ background: #f2f5f7; font-weight: bold; width: 28%; }}
.small {{ font-size: 9pt; }}
</style>
</head>
<body>
<h1>เอกสารฐานข้อมูลเว็บไซต์ JNAC สำหรับสร้างแชทบอท</h1>
<p>รวบรวมจากเว็บไซต์ www.jnac.co.th เพื่อใช้เป็น knowledge base ของธุรกิจ บริษัท เจ แนค (ประเทศไทย) จำกัด</p>
<table class="meta">
<tr><td>แหล่งข้อมูล</td><td>{escape(data['source'])}</td></tr>
<tr><td>วันที่เก็บข้อมูล</td><td>{escape(data['fetched_at'])}</td></tr>
<tr><td>จำนวนหน้าที่ดึงข้อมูล</td><td>{data['page_count']} หน้า</td></tr>
<tr><td>ขอบเขต</td><td>หน้าบริษัท ติดต่อ วิธีสั่งซื้อ Catalogue บทความ และหน้าสินค้า/หมวดสินค้าภายในโดเมน</td></tr>
</table>
<h2>ภาพรวมธุรกิจ</h2>
<ul>
<li>ชื่อบริษัท: บริษัท เจ แนค (ประเทศไทย) จำกัด / J NAC (THAILAND) CO., LTD.</li>
<li>ธุรกิจหลัก: ผลิต จำหน่าย นำเข้า ส่งออก และรับจ้างผลิตสินค้าในกลุ่มอุตสาหกรรมงานขัด ปัด ตัด เจียร รวมถึงงาน ICT Fixture, ICT Tester, Function Test, Assembly Part & Machine</li>
<li>สินค้าหลัก: กระดาษทราย จานทรายซ้อน ใบเจียร ใบตัด ล้อทราย ใยขัดสก๊อตไบร์ท หินขัด ลูกขัด เครื่องมือขัด เครื่องมือลม และสินค้างานขัดแม่พิมพ์</li>
</ul>
<h2>ข้อมูลติดต่อ</h2>
<table class="meta">
<tr><td>ที่อยู่</td><td>84 หมู่ 2 ซ.สุนทรวิภาค ถ.บางพลี-ตำหรุ ต.แพรกษาใหม่ อ.เมือง จ.สมุทรปราการ 10280</td></tr>
<tr><td>โทรศัพท์</td><td>0 2101 5587</td></tr>
<tr><td>มือถือ</td><td>08 0016 1700</td></tr>
<tr><td>แฟกซ์</td><td>0 2101 5587</td></tr>
<tr><td>อีเมล</td><td>info@jnac.co.th, jnac.co.th@gmail.com</td></tr>
<tr><td>Line ID</td><td>@jnac</td></tr>
<tr><td>เลขประจำตัวผู้เสียภาษี</td><td>0115561012346</td></tr>
</table>
<h2>หมวดหมู่ข้อมูล</h2>
<table><tr><th>หมวด</th><th>จำนวนหน้า</th></tr>{category_rows}</table>
<h2>รายการหน้าข้อมูลสำหรับทำแชทบอท</h2>
<table class="small">
<tr><th>#</th><th>หัวข้อ/สินค้า</th><th>หมวด</th><th>ข้อมูลสำคัญที่ดึงได้</th><th>URL</th></tr>
{''.join(rows)}
</table>
</body>
</html>
"""
    OUT_DOC.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    main()
